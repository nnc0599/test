from __future__ import annotations

import base64
from copy import deepcopy
from datetime import datetime
from functools import lru_cache
from html import escape
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.shared import Pt

from app.utils.time_utils import now_utc7


ROOT_DIR = Path(__file__).resolve().parents[2]
TEMPLATE_PATH = ROOT_DIR / "template.docx"
EXPORT_DIR = ROOT_DIR / "data" / "exported_invoices"
PREVIEW_ASSET_DIR = ROOT_DIR / "data" / ".preview_assets"
SCREEN_IMAGE_PATH = ROOT_DIR / "Screen.png"

    
def format_money(value: int) -> str:
    return f"{int(value):,}".replace(",", ".")


def _capitalize_first(text: str) -> str:
    if not text:
        return text
    return text[0].upper() + text[1:]


def _read_three_digits(value: int, full: bool = False) -> str:
    digits = ["không", "một", "hai", "ba", "bốn", "năm", "sáu", "bảy", "tám", "chín"]

    hundred = value // 100
    ten = (value % 100) // 10
    unit = value % 10
    parts: list[str] = []

    if hundred > 0 or full:
        parts.append(f"{digits[hundred]} trăm")

    if ten > 1:
        parts.append(f"{digits[ten]} mươi")
        if unit == 1:
            parts.append("mốt")
        elif unit == 4:
            parts.append("tư")
        elif unit == 5:
            parts.append("lăm")
        elif unit > 0:
            parts.append(digits[unit])
    elif ten == 1:
        parts.append("mười")
        if unit == 5:
            parts.append("lăm")
        elif unit > 0:
            parts.append(digits[unit])
    elif unit > 0:
        if hundred > 0 or full:
            parts.append("lẻ")
        parts.append(digits[unit])

    return " ".join(parts).strip()


def money_to_vietnamese(value: int) -> str:
    if value == 0:
        return "Không đồng./."

    scales = ["", "nghìn", "triệu", "tỷ", "nghìn tỷ", "triệu tỷ"]
    groups: list[int] = []
    number = int(value)
    while number > 0:
        groups.append(number % 1000)
        number //= 1000

    parts: list[str] = []
    highest_index = len(groups) - 1
    for index in range(highest_index, -1, -1):
        group = groups[index]
        if group == 0:
            continue
        full = index < highest_index and groups[index + 1] > 0 and group < 100
        text = _read_three_digits(group, full=full)
        scale = scales[index] if index < len(scales) else ""
        parts.append(f"{text} {scale}".strip())

    return f"{_capitalize_first(' '.join(parts).strip())} đồng./."


def _parse_created_at(created_at: str) -> datetime:
    normalized = created_at.strip().replace("T", " ")
    for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%d"):
        try:
            return datetime.strptime(normalized, fmt)
        except ValueError:
            continue
    raise ValueError(f"Không đọc được ngày tạo hóa đơn: {created_at}")


def _set_paragraph_text(paragraph, text: str, alignment=None, bold: bool | None = None, italic: bool | None = None) -> None:
    paragraph.text = text
    if alignment is not None:
        paragraph.alignment = alignment
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        _set_run_font(run, bold=bold, italic=italic)


def _set_run_font(run, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ["w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"]:
        r_fonts.set(qn(attr), "Times New Roman")
    size = r_pr.find(qn("w:sz"))
    if size is None:
        size = OxmlElement("w:sz")
        r_pr.append(size)
    size.set(qn("w:val"), "24")
    size_cs = r_pr.find(qn("w:szCs"))
    if size_cs is None:
        size_cs = OxmlElement("w:szCs")
        r_pr.append(size_cs)
    size_cs.set(qn("w:val"), "24")


def _set_cell_text(
    cell,
    text: str,
    alignment=None,
    bold: bool | None = None,
    italic: bool | None = None,
    vertical_alignment=None,
) -> None:
    cell.text = text
    if vertical_alignment is not None:
        cell.vertical_alignment = vertical_alignment
    for paragraph in cell.paragraphs:
        if alignment is not None:
            paragraph.alignment = alignment
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            _set_run_font(run, bold=bold, italic=italic)


def _set_paragraph_segments(paragraph, segments: list[tuple[str, bool | None, bool | None]], alignment=None) -> None:
    paragraph.clear()
    if alignment is not None:
        paragraph.alignment = alignment
    paragraph.paragraph_format.space_after = Pt(0)
    for text, bold, italic in segments:
        run = paragraph.add_run(text)
        _set_run_font(run, bold=bold, italic=italic)


def _set_cell_segments(cell, segments: list[tuple[str, bool | None, bool | None]], alignment=None) -> None:
    if not cell.paragraphs:
        paragraph = cell.add_paragraph()
    else:
        paragraph = cell.paragraphs[0]
    paragraph.clear()
    if alignment is not None:
        paragraph.alignment = alignment
    paragraph.paragraph_format.space_after = Pt(0)
    for extra_paragraph in cell.paragraphs[1:]:
        _remove_paragraph(extra_paragraph)
    for text, bold, italic in segments:
        run = paragraph.add_run(text)
        _set_run_font(run, bold=bold, italic=italic)


def _set_table_row_height(table, height_cm: float) -> None:
    row_height = Cm(height_cm)
    for row in table.rows:
        row.height = row_height
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY


def _remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _set_table_no_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)

    for edge_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        edge = tbl_borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            tbl_borders.append(edge)
        edge.set(qn("w:val"), "nil")


def _set_table_grid_borders(table, color: str = "9CA3AF", size: str = "6") -> None:
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)

    for edge_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        edge = tbl_borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            tbl_borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), size)
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), color)


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_document_style(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)


def _add_spacer(doc: Document, size_pt: int = 4) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("")
    _set_run_font(run)
    run.font.size = Pt(size_pt)


def _add_report_title(doc: Document, title: str, subtitle: str | None = None) -> None:
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(title)
    _set_run_font(title_run, bold=True)
    title_run.font.size = Pt(16)

    if subtitle:
        subtitle_paragraph = doc.add_paragraph()
        subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_paragraph.add_run(subtitle)
        _set_run_font(subtitle_run, italic=True)
        subtitle_run.font.size = Pt(11)


def _add_report_meta_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_grid_borders(table, color="CBD5E1")
    _set_table_row_height(table, 0.7)

    widths = [Cm(4.2), Cm(12.0)]
    for row in table.rows:
        row.cells[0].width = widths[0]
        row.cells[1].width = widths[1]

    for row_index, (label, value) in enumerate(rows):
        _shade_cell(table.cell(row_index, 0), "E2E8F0")
        _set_cell_text(
            table.cell(row_index, 0),
            label,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            bold=True,
            vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER,
        )
        _set_cell_text(
            table.cell(row_index, 1),
            value,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER,
        )


def _add_report_data_table(
    doc: Document,
    headers: list[str],
    data_rows: list[list[str]],
    column_widths_cm: list[float],
    centered_columns: set[int] | None = None,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_grid_borders(table, color="94A3B8")
    _set_table_row_height(table, 0.7)

    for index, width in enumerate(column_widths_cm):
        table.rows[0].cells[index].width = Cm(width)

    for index, header in enumerate(headers):
        header_cell = table.rows[0].cells[index]
        _shade_cell(header_cell, "DCE6F1")
        _set_cell_text(
            header_cell,
            header,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            bold=True,
            vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER,
        )

    if not data_rows:
        row = table.add_row()
        row.height = Cm(0.7)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        empty_cell = row.cells[0]
        if len(headers) > 1:
            empty_cell = row.cells[0].merge(row.cells[len(headers) - 1])
        _set_cell_text(
            empty_cell,
            "Không có dữ liệu trong kỳ báo cáo đã chọn.",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            italic=True,
            vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER,
        )
        return

    for values in data_rows:
        row = table.add_row()
        row.height = Cm(0.7)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for col_index, value in enumerate(values):
            _set_cell_text(
                row.cells[col_index],
                value,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER,
            )


def _add_signature_section(doc: Document, generated_at_text: str) -> None:
    location_paragraph = doc.add_paragraph()
    location_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    location_run = location_paragraph.add_run(generated_at_text)
    _set_run_font(location_run, italic=True)

    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_no_borders(table)

    for row in table.rows:
        row.cells[0].width = Cm(8.0)
        row.cells[1].width = Cm(8.0)

    _set_cell_text(table.cell(0, 0), "Người làm báo cáo", alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    _set_cell_text(table.cell(0, 1), "Giám đốc", alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    _set_cell_text(table.cell(1, 0), "(Ký, ghi rõ họ tên)", alignment=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    _set_cell_text(table.cell(1, 1), "(Ký, ghi rõ họ tên)", alignment=WD_ALIGN_PARAGRAPH.CENTER, italic=True)

    for _ in range(4):
        _add_spacer(doc, size_pt=6)


def export_report_docx(
    *,
    report_title: str,
    period_label: str,
    summary_rows: list[tuple[str, str]],
    headers: list[str],
    data_rows: list[list[str]],
    output_name: str,
    column_widths_cm: list[float],
    centered_columns: set[int] | None = None,
    output_dir: Path | None = None,
) -> Path:
    doc = Document()
    _set_document_style(doc)

    generated_at = now_utc7()
    generated_at_label = generated_at.strftime("Ngày %d tháng %m năm %Y, %H:%M:%S")
    subtitle = f"Kỳ báo cáo: {period_label}" if period_label else None

    _add_report_title(doc, report_title, subtitle)
    _add_spacer(doc)
    _add_report_meta_table(
        doc,
        [
            ("Thời gian tạo báo cáo", generated_at.strftime("%d/%m/%Y %H:%M:%S")),
            *summary_rows,
        ],
    )
    _add_spacer(doc)
    _add_report_data_table(doc, headers, data_rows, column_widths_cm, centered_columns)
    _add_spacer(doc)
    _add_signature_section(doc, generated_at_label)

    save_dir = output_dir or EXPORT_DIR
    save_dir.mkdir(parents=True, exist_ok=True)
    output_path = save_dir / output_name
    doc.save(str(output_path))
    return output_path


def _build_customer_info_table(doc: Document, invoice: dict):
    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_no_borders(table)

    column_widths = [Cm(10.4), Cm(5.2)]
    for row in table.rows:
        for index, width in enumerate(column_widths):
            row.cells[index].width = width

    _set_cell_segments(
        table.cell(0, 0),
        [("Khách hàng: ", True, None), (invoice.get('customer_name', '').strip(), None, None)],
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _set_cell_segments(
        table.cell(0, 1),
        [("Điện thoại: ", True, None), (invoice.get('phone', '').strip(), None, None)],
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _set_cell_segments(
        table.cell(1, 0),
        [("Gmail: ", True, None), (invoice.get('email', '').strip(), None, None)],
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _set_cell_segments(
        table.cell(1, 1),
        [("Mã số thuế: ", True, None), (invoice.get('tax_code', '').strip(), None, None)],
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    address_cell = table.cell(2, 0).merge(table.cell(2, 1))
    _set_cell_text(
        address_cell,
        f"Địa chỉ: {invoice.get('address', '').strip()}",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    return table


def _populate_item_row(row, index: int, item: dict) -> None:
    values = [
        str(index),
        item.get("product_code", ""),
        item.get("product_name", ""),
        item.get("unit", ""),
        str(item.get("quantity", "")),
        format_money(int(item.get("unit_price", 0))),
        format_money(int(item.get("line_total", 0))),
    ]
    centered_columns = {0, 3, 4, 5, 6}
    for column_index, (cell, value) in enumerate(zip(row.cells, values)):
        alignment = WD_ALIGN_PARAGRAPH.CENTER if column_index in centered_columns else None
        _set_cell_text(cell, value, alignment=alignment)


def _image_file_to_data_uri(path: Path) -> str:
    if not path.exists() or not path.is_file():
        return ""

    suffix = path.suffix.lower()
    mime_type = {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".webp": "image/webp",
    }.get(suffix, "application/octet-stream")
    encoded = base64.b64encode(path.read_bytes()).decode("ascii")
    return f"data:{mime_type};base64,{encoded}"


@lru_cache(maxsize=1)
def _load_template_preview_assets() -> dict:
    company_lines: list[str] = []
    account_label = "Tài khoản"
    logo_uri = ""
    qr_uri = ""
    page_margin_mm = {"top": 25.0, "right": 15.0, "bottom": 20.0, "left": 23.0}
    header_widths = [20.5, 61.3, 18.2]
    item_widths = [4.75, 21.54, 35.91, 6.57, 5.80, 11.98, 14.69]
    logo_size_mm = {"width": 31.5, "height": 23.7}
    qr_size_mm = {"width": 25.9, "height": 21.6}

    if TEMPLATE_PATH.exists():
        doc = Document(str(TEMPLATE_PATH))
        if doc.tables:
            header_table = doc.tables[0]
            if len(header_table.rows) > 0:
                company_lines = [
                    line.strip()
                    for line in header_table.cell(0, 1).text.splitlines()
                    if line.strip()
                ]
                account_label = header_table.cell(0, 2).text.strip() or account_label

        with ZipFile(TEMPLATE_PATH) as archive:
            media_names = sorted(
                name
                for name in archive.namelist()
                if name.startswith("word/media/") and not name.endswith("/")
            )
            PREVIEW_ASSET_DIR.mkdir(parents=True, exist_ok=True)
            if media_names:
                logo_path = PREVIEW_ASSET_DIR / Path(media_names[0]).name
                logo_path.write_bytes(archive.read(media_names[0]))
                logo_uri = logo_path.name
            if len(media_names) > 1:
                qr_path = PREVIEW_ASSET_DIR / Path(media_names[1]).name
                qr_path.write_bytes(archive.read(media_names[1]))
                qr_uri = qr_path.name

            from xml.etree import ElementTree as ET

            ns = {
                "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
                "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
            }
            root = ET.fromstring(archive.read("word/document.xml"))

            def twips_to_mm(value: str) -> float:
                return round(int(value) * 25.4 / 1440, 2)

            def emu_to_mm(value: str) -> float:
                return round(int(value) / 36000, 2)

            sect = root.find('.//w:sectPr', ns)
            if sect is not None:
                margins = sect.find('w:pgMar', ns)
                if margins is not None:
                    page_margin_mm = {
                        "top": twips_to_mm(margins.attrib.get(f"{{{ns['w']}}}top", "1417")),
                        "right": twips_to_mm(margins.attrib.get(f"{{{ns['w']}}}right", "850")),
                        "bottom": twips_to_mm(margins.attrib.get(f"{{{ns['w']}}}bottom", "1134")),
                        "left": twips_to_mm(margins.attrib.get(f"{{{ns['w']}}}left", "1304")),
                    }

            tables = root.findall('.//w:tbl', ns)
            if len(tables) > 0:
                grid = tables[0].find('./w:tblGrid', ns)
                if grid is not None:
                    cols = [int(col.attrib[f"{{{ns['w']}}}w"]) for col in grid.findall('w:gridCol', ns)]
                    total = sum(cols) or 1
                    header_widths = [round(col * 100 / total, 2) for col in cols]

                drawings = tables[0].findall('.//w:drawing', ns)
                if len(drawings) > 0:
                    extent = drawings[0].find('.//wp:extent', ns)
                    if extent is not None:
                        logo_size_mm = {
                            "width": emu_to_mm(extent.attrib.get("cx", "1188085")),
                            "height": emu_to_mm(extent.attrib.get("cy", "896620")),
                        }
                if len(drawings) > 1:
                    extent = drawings[1].find('.//wp:extent', ns)
                    if extent is not None:
                        qr_size_mm = {
                            "width": emu_to_mm(extent.attrib.get("cx", "977900")),
                            "height": emu_to_mm(extent.attrib.get("cy", "817245")),
                        }

            if len(tables) > 1:
                grid = tables[1].find('./w:tblGrid', ns)
                if grid is not None:
                    cols = [int(col.attrib[f"{{{ns['w']}}}w"]) for col in grid.findall('w:gridCol', ns)]
                    total = sum(cols) or 1
                    item_widths = [round(col * 100 / total, 2) for col in cols]

    return {
        "company_lines": company_lines,
        "account_label": account_label,
        "logo_uri": logo_uri,
        "qr_uri": qr_uri,
        "page_margin_mm": page_margin_mm,
        "header_widths": header_widths,
        "item_widths": item_widths,
        "logo_size_mm": logo_size_mm,
        "qr_size_mm": qr_size_mm,
    }


def build_invoice_preview_html(invoice: dict, items: list[dict]) -> str:
    created_at = _parse_created_at(str(invoice.get("created_at", "")))
    assets = _load_template_preview_assets()
    page_margin_mm = assets["page_margin_mm"]
    item_widths = assets["item_widths"]
    screen_image_uri = _image_file_to_data_uri(SCREEN_IMAGE_PATH)
    header_html = (
        f"<div class=\"header-image-wrap\"><img class=\"header-image\" src=\"{screen_image_uri}\" /></div>"
        if screen_image_uri
        else ""
    )

    item_rows = "".join(
        f"""
        <tr>
            <td class=\"col-stt\">{index}</td>
            <td class=\"col-code\">{escape(str(item.get('product_code', '')))}</td>
            <td class=\"col-name\">{escape(str(item.get('product_name', '')))}</td>
            <td class=\"col-unit\">{escape(str(item.get('unit', '')))}</td>
            <td class=\"col-qty\">{escape(str(item.get('quantity', '')))}</td>
            <td class=\"col-price\">{format_money(int(item.get('unit_price', 0)))}</td>
            <td class=\"col-total\">{format_money(int(item.get('line_total', 0)))}</td>
        </tr>
        """
        for index, item in enumerate(items, start=1)
    )

    total_amount = int(invoice.get("total_amount", 0))
    return f"""
        <html>
            <head>
                <style>
                    body {{
                        font-family: 'Times New Roman';
                        font-size: 12pt;
                        color: #111827;
                        line-height: 1.25;
                        background: #ffffff;
                        margin: 0;
                    }}
                    .sheet {{
                        width: 210mm;
                        min-height: 297mm;
                        box-sizing: border-box;
                        padding: {page_margin_mm['top']}mm {page_margin_mm['right']}mm {page_margin_mm['bottom']}mm {page_margin_mm['left']}mm;
                        text-align: center;
                    }}
                    .header {{
                        width: 100%;
                        margin-bottom: 3.2mm;
                    }}
                    .header-image-wrap {{
                        width: 100%;
                        display: block;
                        margin: 0 auto;
                        text-align: center;
                    }}
                    .header-image {{
                        width: 100%;
                        max-width: 100%;
                        height: auto;
                        display: block;
                        margin: 0 auto;
                    }}
                    .title {{
                        text-align: center;
                        font-size: 19px;
                        font-weight: bold;
                        letter-spacing: 0;
                        margin: 4mm 0 3mm;
                    }}
                    .meta-table, .info-table, .items-table, .signatures {{
                        width: 100%;
                        border-collapse: collapse;
                    }}
                    .meta-table td, .info-table td, .signatures td {{
                        border: none;
                        padding: 0;
                    }}
                    .meta-left {{
                        text-align: center;
                    }}
                    .meta-table {{
                        margin: 0 0 1.5mm;
                    }}
                    .meta-table td {{
                        padding: 0;
                    }}
                    .meta-invoice-row td {{
                        padding-top: 0.8mm;
                        font-style: italic;
                    }}
                    .info-table {{
                        margin-bottom: 2.5mm;
                        table-layout: fixed;
                    }}
                    .info-table td {{
                        padding: 0.6mm 0;
                        text-align: center;
                        vertical-align: top;
                    }}
                    .info-left {{
                        width: 66.67%;
                        padding-right: 4mm;
                    }}
                    .info-right {{
                        width: 33.33%;
                        text-align: center;
                    }}
                    .info-address {{
                        padding-top: 1mm;
                    }}
                    .items-table {{
                        margin-top: 0.8mm;
                    }}
                    .items-table th, .items-table td {{
                        border: 1px solid #111827;
                        padding: 1.2mm 0.9mm;
                        text-align: left;
                        vertical-align: middle;
                        font-size: 12pt;
                    }}
                    .items-table th {{
                        font-weight: bold;
                    }}
                    .items-table .col-stt,
                    .items-table .col-unit,
                    .items-table .col-qty,
                    .items-table .col-price,
                    .items-table .col-total {{
                        text-align: center;
                    }}
                    .items-table .col-name {{
                        text-align: center;
                        padding-left: 0.9mm;
                    }}
                    .col-stt {{ width: {item_widths[0]}%; }}
                    .col-code {{ width: {item_widths[1]}%; }}
                    .col-name {{ width: {item_widths[2]}%; }}
                    .col-unit {{ width: {item_widths[3]}%; }}
                    .col-qty {{ width: {item_widths[4]}%; }}
                    .col-price {{ width: {item_widths[5]}%; }}
                    .col-total {{ width: {item_widths[6]}%; }}
                    .total-label {{
                        text-align: center;
                        font-weight: bold;
                    }}
                    .money-words {{
                        margin-top: 1.8mm;
                        font-size: 12pt;
                        text-align: center;
                        font-style: italic;
                    }}
                    .total-row td {{
                        font-weight: bold;
                    }}
                    .signatures {{
                        margin-top: 14mm;
                    }}
                    .signatures td {{
                        width: 25%;
                        text-align: center;
                        padding-top: 2mm;
                        font-size: 11px;
                    }}
                    .sign-label {{
                        font-weight: bold;
                    }}
                    .sign-note {{
                        font-style: italic;
                        margin-top: 3px;
                    }}
                </style>
            </head>
            <body>
                <div class="sheet">
                    <div class="header">{header_html}</div>

                    <div class="title">PHIẾU XUẤT KHO KIÊM BẢO HÀNH</div>

                    <table class="meta-table">
                        <tr>
                            <td class="meta-left">Ngày {created_at.day:02d} tháng {created_at.month:02d} năm {created_at.year}</td>
                        </tr>
                        <tr class="meta-invoice-row">
                            <td class="meta-left"><strong>Số:</strong> {escape(str(invoice.get('invoice_no', '')))}</td>
                        </tr>
                    </table>

                    <table class="info-table">
                        <tr>
                            <td class="info-left"><strong>Khách hàng:</strong> {escape(str(invoice.get('customer_name', '')))}</td>
                            <td class="info-right"><strong>Điện thoại:</strong> {escape(str(invoice.get('phone', '')))}</td>
                        </tr>
                        <tr>
                            <td class="info-left"><strong>Gmail:</strong> {escape(str(invoice.get('email', '')))}</td>
                            <td class="info-right"><strong>Mã số thuế:</strong> {escape(str(invoice.get('tax_code', '')))}</td>
                        </tr>
                        <tr>
                            <td class="info-address" colspan="2"><strong>Địa chỉ:</strong> {escape(str(invoice.get('address', '')))}</td>
                        </tr>
                    </table>

                    <table class="items-table">
                        <thead>
                            <tr>
                                <th class="col-stt">TT</th>
                                <th class="col-code">Mã hàng</th>
                                <th class="col-name">Tên sản phẩm</th>
                                <th class="col-unit">ĐV</th>
                                <th class="col-qty">SL</th>
                                <th class="col-price">Đơn giá</th>
                                <th class="col-total">Thành tiền</th>
                            </tr>
                        </thead>
                        <tbody>
                            {item_rows}
                            <tr class="total-row">
                                <td class="total-label" colspan="3">Tổng thanh toán</td>
                                <td></td>
                                <td></td>
                                <td></td>
                                <td>{format_money(total_amount)}</td>
                            </tr>
                        </tbody>
                    </table>

                    <div class="money-words"><strong>Số tiền viết bằng chữ:</strong> {escape(money_to_vietnamese(total_amount))}</div>

                    <table class="signatures">
                        <tr>
                            <td>
                                <div class="sign-label">Khách hàng</div>
                                <div class="sign-note">(Ký, họ tên)</div>
                            </td>
                            <td>
                                <div class="sign-label">Người lập phiếu</div>
                                <div class="sign-note">(Ký, họ tên)</div>
                            </td>
                            <td>
                                <div class="sign-label">Thủ kho</div>
                                <div class="sign-note">(Ký, họ tên)</div>
                            </td>
                            <td>
                                <div class="sign-label">Giám đốc</div>
                                <div class="sign-note">(Ký, họ tên, đóng dấu)</div>
                            </td>
                        </tr>
                    </table>
                </div>
            </body>
        </html>
        """


def export_invoice_docx(invoice: dict, items: list[dict], output_dir: Path | None = None) -> Path:
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Không tìm thấy mẫu hóa đơn: {TEMPLATE_PATH}")
    if not items:
        raise ValueError("Không có sản phẩm để xuất file hóa đơn")

    doc = Document(str(TEMPLATE_PATH))
    created_at = _parse_created_at(str(invoice.get("created_at", "")))
    invoice_no = str(invoice.get("invoice_no", "")).strip()

    if len(doc.paragraphs) < 7 or len(doc.tables) < 2:
        raise ValueError("Mẫu template.docx không đúng cấu trúc mong đợi")

    items_table = doc.tables[1]

    date_paragraph = doc.paragraphs[1]
    invoice_paragraph = doc.paragraphs[2]
    email_paragraph = doc.paragraphs[3]
    address_paragraph = doc.paragraphs[4]
    money_words_paragraph = doc.paragraphs[6]

    _set_paragraph_text(
        date_paragraph,
        f"Ngày {created_at.day:02d} tháng {created_at.month:02d} năm {created_at.year}",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _set_paragraph_segments(
        invoice_paragraph,
        [("Số: ", True, True), (invoice_no, None, True)],
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    customer_table = _build_customer_info_table(doc, invoice)
    invoice_paragraph._p.addnext(customer_table._tbl)
    _remove_paragraph(email_paragraph)
    _remove_paragraph(address_paragraph)

    _set_paragraph_text(
        money_words_paragraph,
        f"Số tiền viết bằng chữ: {money_to_vietnamese(int(invoice.get('total_amount', 0)))}",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        italic=True,
    )

    for column_index in [0, 3, 4, 5, 6]:
        for paragraph in items_table.rows[0].cells[column_index].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(items_table.rows) < 3:
        raise ValueError("Bảng hàng hóa trong template.docx không đúng cấu trúc mong đợi")

    template_tr = deepcopy(items_table.rows[1]._tr)
    total_tr = items_table.rows[2]._tr
    for _ in range(len(items) - 1):
        items_table._tbl.insert(items_table._tbl.index(total_tr), deepcopy(template_tr))

    for index, item in enumerate(items, start=1):
        _populate_item_row(items_table.rows[index], index, item)

    total_row = items_table.rows[len(items) + 1]
    _set_cell_text(total_row.cells[0], "Tổng thanh toán", alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    _set_cell_text(total_row.cells[6], format_money(int(invoice.get("total_amount", 0))), alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

    save_dir = output_dir or EXPORT_DIR
    save_dir.mkdir(parents=True, exist_ok=True)
    output_path = save_dir / f"{invoice_no}.docx"
    doc.save(str(output_path))
    return output_path